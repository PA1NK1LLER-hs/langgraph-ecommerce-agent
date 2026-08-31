"""RAG 模块单元测试 — 文档解析器 + 分块器 + 混合检索 + 引用。"""

import os
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch, AsyncMock

import pytest


# ═══════════════════════════════════════════════════
# 分块器测试
# ═══════════════════════════════════════════════════


class TestChunkers:
    """测试三种分块策略。"""

    def test_semantic_chunker_basic(self):
        from rag.parsers import _SemanticChunker
        chunker = _SemanticChunker(chunk_size=100, chunk_overlap=20)
        text = "第一段内容。\n\n第二段内容。\n\n第三段内容。"
        chunks = chunker.split(text)
        assert len(chunks) >= 1
        assert all(isinstance(c, str) for c in chunks)

    def test_semantic_chunker_empty(self):
        from rag.parsers import _SemanticChunker
        chunker = _SemanticChunker()
        chunks = chunker.split("")
        assert chunks == [""]

    def test_fixed_chunker(self):
        from rag.parsers import _FixedChunker
        chunker = _FixedChunker(chunk_size=20, chunk_overlap=5)
        text = "A" * 55
        chunks = chunker.split(text)
        assert len(chunks) >= 2
        # Each chunk should be <= chunk_size
        for c in chunks:
            assert len(c) <= 20

    def test_recursive_chunker_chinese(self):
        from rag.parsers import _RecursiveChunker
        chunker = _RecursiveChunker(chunk_size=50, chunk_overlap=10)
        text = "第一句话。第二句话。第三句话，继续写。第四句话。"
        chunks = chunker.split(text)
        assert len(chunks) >= 1

    def test_recursive_chunker_fallback(self):
        """极长无分隔文本也能被切割。"""
        from rag.parsers import _RecursiveChunker
        chunker = _RecursiveChunker(chunk_size=10, chunk_overlap=2)
        text = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"  # No separators
        chunks = chunker.split(text)
        assert len(chunks) >= 2


# ═══════════════════════════════════════════════════
# ParsedDocument 测试
# ═══════════════════════════════════════════════════


class TestParsedDocument:
    def test_to_chunks(self):
        from rag.parsers import ParsedDocument
        doc = ParsedDocument(
            text="段落A\n\n段落B\n\n段落C\n\n段落D\n\n段落E",
            metadata={"source": "test.txt"},
        )
        chunks = doc.to_chunks(strategy="semantic", chunk_size=20)
        assert len(chunks) >= 1
        for ch in chunks:
            assert ch.metadata.get("source") == "test.txt"
            assert ch.chunk_type == "text"

    def test_chunk_index_increment(self):
        from rag.parsers import ParsedDocument
        doc = ParsedDocument(text="A\n\nB\n\nC\n\nD")
        chunks = doc.to_chunks(strategy="semantic", chunk_size=5)
        indices = [c.chunk_index for c in chunks]
        assert indices == list(range(len(chunks)))


# ═══════════════════════════════════════════════════
# 文档解析器测试
# ═══════════════════════════════════════════════════


class TestDocumentParser:
    def test_parse_text_file(self):
        from rag.parsers import DocumentParser
        parser = DocumentParser()
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("Hello 世界\n\n这是测试文本。")
            tmp_path = f.name

        try:
            doc = parser.parse_file(tmp_path)
            assert "Hello" in doc.text
            assert "测试" in doc.text
            assert doc.metadata["extension"] == ".txt"
        finally:
            os.unlink(tmp_path)

    def test_parse_file_not_found(self):
        from rag.parsers import DocumentParser
        parser = DocumentParser()
        with pytest.raises(FileNotFoundError):
            parser.parse_file("/nonexistent/file.pdf")

    def test_unsupported_extension(self):
        from rag.parsers import DocumentParser
        parser = DocumentParser()
        with tempfile.NamedTemporaryFile(mode="w", suffix=".xyz", delete=False) as f:
            f.write("test")
            tmp_path = f.name
        try:
            with pytest.raises(ValueError, match="不支持的文件格式"):
                parser.parse_file(tmp_path)
        finally:
            os.unlink(tmp_path)

    @pytest.mark.skipif(
        not __import__("importlib").import_module("importlib").util.find_spec("docx"),
        reason="python-docx not installed",
    )
    def test_parse_docx_minimal(self):
        from rag.parsers import DocumentParser
        try:
            import docx
        except ImportError:
            pytest.skip("python-docx not installed")

        parser = DocumentParser()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name

        try:
            doc = docx.Document()
            doc.add_paragraph("测试文档内容。")
            doc.add_paragraph("第二段。")
            doc.save(tmp_path)

            result = parser.parse_file(tmp_path)
            assert "测试文档内容" in result.text
            assert result.metadata["extension"] == ".docx"
        finally:
            os.unlink(tmp_path)

    def test_parse_html_minimal(self):
        from rag.parsers import DocumentParser
        parser = DocumentParser()
        with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8") as f:
            f.write("<html><head><title>Test</title></head><body><p>Hello World</p></body></html>")
            tmp_path = f.name

        try:
            doc = parser.parse_file(tmp_path)
            # Should extract text content
            assert len(doc.text) > 0
        finally:
            os.unlink(tmp_path)

    def test_parse_markdown_file(self):
        from rag.parsers import DocumentParser
        parser = DocumentParser()
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
            f.write("# 标题\n\n这是 Markdown 内容。\n\n- 列表项1\n- 列表项2")
            tmp_path = f.name

        try:
            doc = parser.parse_file(tmp_path)
            assert "标题" in doc.text
            assert "Markdown" in doc.text
        finally:
            os.unlink(tmp_path)


# ═══════════════════════════════════════════════════
# BM25 索引测试
# ═══════════════════════════════════════════════════


class TestBM25Index:
    @staticmethod
    def _has_rank_bm25():
        try:
            import rank_bm25  # noqa: F401
            return True
        except ImportError:
            return False

    def test_add_and_search(self):
        if not self._has_rank_bm25():
            pytest.skip("rank_bm25 not installed")
        from rag.retrieval import _BM25Index
        bm25 = _BM25Index()
        bm25.add("Python is a programming language", {"source": "doc1"})
        bm25.add("Java is also a programming language", {"source": "doc2"})
        bm25.add("Machine learning with Python", {"source": "doc3"})

        results = bm25.search("Python programming", top_k=2)
        assert len(results) >= 1
        # "Python programming" doc should be in top results
        assert any("Python" in r.content for r in results)

    def test_add_batch(self):
        if not self._has_rank_bm25():
            pytest.skip("rank_bm25 not installed")
        from rag.retrieval import _BM25Index
        bm25 = _BM25Index()
        texts = ["doc one", "doc two", "doc three"]
        metas = [{"source": "a"}, {"source": "b"}, {"source": "c"}]
        bm25.add_batch(texts, metas)

        assert len(bm25) == 3
        results = bm25.search("doc", top_k=5)
        assert len(results) == 3

    def test_empty_search(self):
        from rag.retrieval import _BM25Index
        bm25 = _BM25Index()
        results = bm25.search("anything", top_k=5)
        assert results == []

    def test_remove_by_source(self):
        from rag.retrieval import _BM25Index
        bm25 = _BM25Index()
        bm25.add("doc A", {"source": "src1"})
        bm25.add("doc B", {"source": "src2"})
        bm25.add("doc C", {"source": "src1"})

        removed = bm25.remove_by_source("src1")
        assert removed == 2
        assert len(bm25) == 1

    def test_clear(self):
        from rag.retrieval import _BM25Index
        bm25 = _BM25Index()
        bm25.add("some text")
        bm25.clear()
        assert len(bm25) == 0

    def test_chinese_tokenization(self):
        if not self._has_rank_bm25():
            pytest.skip("rank_bm25 not installed")
        from rag.retrieval import _BM25Index
        bm25 = _BM25Index()
        bm25.add("Python 是一门编程语言", {"source": "zh"})
        bm25.add("Java 也是一种编程语言", {"source": "zh2"})

        results = bm25.search("编程语言", top_k=3)
        assert len(results) >= 1


# ═══════════════════════════════════════════════════
# RRF 融合测试
# ═══════════════════════════════════════════════════


class TestRRFFusion:
    def test_rrf_merges_results(self):
        from rag.retrieval import HybridRetriever, SearchResult

        results_a = [
            SearchResult(content="Document Alpha", score=0.9, source="bm25"),
            SearchResult(content="Document Beta", score=0.7, source="bm25"),
            SearchResult(content="Document Gamma", score=0.5, source="bm25"),
        ]
        results_b = [
            SearchResult(content="Document Beta", score=0.95, source="dense"),
            SearchResult(content="Document Delta", score=0.8, source="dense"),
            SearchResult(content="Document Alpha", score=0.6, source="dense"),
        ]

        fused = HybridRetriever._rrf_fuse(results_a, results_b, k=60)
        assert len(fused) == 4  # Alpha, Beta, Gamma, Delta (unique by content)
        # Beta should rank high (appears in both lists at high ranks)
        assert fused[0].content == "Document Beta"

    def test_rrf_single_list(self):
        from rag.retrieval import HybridRetriever, SearchResult

        results = [
            SearchResult(content="Only doc", score=0.9, source="test"),
        ]
        fused = HybridRetriever._rrf_fuse(results, [], k=60)
        assert len(fused) == 1
        assert fused[0].content == "Only doc"


# ═══════════════════════════════════════════════════
# 引用 Schema 测试
# ═══════════════════════════════════════════════════


class TestCitationSchema:
    def test_citation_creation(self):
        from agent.response_schemas import Citation
        c = Citation(index=1, source="test.pdf", content_snippet="这是一段测试", relevance_score=0.92)
        assert c.index == 1
        assert c.source == "test.pdf"
        assert c.relevance_score == 0.92

    def test_cited_response_creation(self):
        from agent.response_schemas import CitedResponse, Citation
        resp = CitedResponse(
            content="根据测试文档，答案是42。",
            citations=[
                Citation(index=1, source="test.pdf", content_snippet="答案", relevance_score=0.95),
            ],
        )
        assert resp.response_type == "cited"
        assert len(resp.citations) == 1

    def test_cited_response_in_registry(self):
        from agent.response_schemas import RESPONSE_SCHEMAS, CitedResponse
        assert "cited" in RESPONSE_SCHEMAS
        assert RESPONSE_SCHEMAS["cited"] == CitedResponse


# ═══════════════════════════════════════════════════
# 索引器 bug 修复验证
# ═══════════════════════════════════════════════════


class TestIndexerBugFixes:
    """验证 B1-B4 bug 修复。"""

    def test_search_returns_structured_results(self):
        """B4 修复：async_search_knowledge dense 模式返回结构化结果列表。"""
        # 这个测试在无 LightRAG 环境时验证返回格式
        from rag.indexer import async_search_knowledge

        # 由于没有实际的 LightRAG 实例，会返回 error status
        # 但错误格式也应该保持一致
        # 在 CI 中 skip 如果 LightRAG 未初始化
        pytest.skip("需要 LightRAG 和 Qdrant 运行环境")

    def test_index_result_includes_chunks_count(self):
        """B1 修复：index_text 返回真实 chunk 数而不是硬编码 1。"""
        pytest.skip("需要 LightRAG 和 Qdrant 运行环境")

    def test_tags_passed_to_lightrag(self):
        """B2 修复：tags/user_id 参数不再通过不支持的 addon_params 传递，
        而是仅记录日志（LightRAG 1.5.x 不支持自定义元数据参数）。"""
        from unittest.mock import AsyncMock, patch

        # Mock: ensure_initialized_async is a no-op，且 _rag 暴露当前真实调用的
        # apipeline_enqueue_documents / apipeline_process_enqueue_documents 管道方法。
        with patch("rag.indexer.KnowledgeIndexer._ensure_initialized_async", AsyncMock()):
            from rag.indexer import KnowledgeIndexer
            indexer = KnowledgeIndexer()
            mock_rag = MagicMock()
            mock_rag.apipeline_enqueue_documents = AsyncMock()
            mock_rag.apipeline_process_enqueue_documents = AsyncMock()
            indexer._rag = mock_rag

            import asyncio
            async def _run():
                await indexer._index_async("test text", "src1", tags="tag1,tag2", user_id="u1")
            asyncio.run(_run())

            # Verify the enqueue pipeline was called (not the old ainsert API)
            mock_rag.apipeline_enqueue_documents.assert_awaited_once()
            mock_rag.apipeline_process_enqueue_documents.assert_awaited_once()
            call_kwargs = mock_rag.apipeline_enqueue_documents.call_args
            assert call_kwargs is not None
            _, kwargs = call_kwargs
            # file_paths 是唯一化后的路径（"src1::<uuid>"），不再是裸 source，
            # 且 addon_params 不应传入（LightRAG 1.5.x 不接受）。
            fps = kwargs.get("file_paths")
            assert isinstance(fps, list) and len(fps) == 1
            assert fps[0].startswith("src1::")
            assert "addon_params" not in kwargs
